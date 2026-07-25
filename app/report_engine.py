from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.enums import TA_RIGHT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

try:
    import arabic_reshaper
    from bidi.algorithm import get_display
except ImportError:  # pragma: no cover - dependencies installed from requirements.txt
    arabic_reshaper = None
    get_display = None

FONT_PATHS = (
    Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    Path("C:/Windows/Fonts/arial.ttf"),
)


def _rtl(text: object) -> str:
    value = "" if text is None else str(text)
    if arabic_reshaper and get_display:
        return get_display(arabic_reshaper.reshape(value))
    return value


def _register_pdf_font() -> str:
    for path in FONT_PATHS:
        if path.exists():
            name = "PharmaArabic"
            if name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(name, str(path)))
            return name
    return "Helvetica"


def build_pdf(title: str, company_name: str, subtitle: str, columns: Sequence[str], rows: Iterable[Sequence[object]], metadata: Sequence[str]) -> BytesIO:
    output = BytesIO()
    font_name = _register_pdf_font()
    doc = SimpleDocTemplate(
        output,
        pagesize=landscape(A4),
        rightMargin=12 * mm,
        leftMargin=12 * mm,
        topMargin=12 * mm,
        bottomMargin=12 * mm,
        title=title,
        author=company_name,
    )
    right = ParagraphStyle("rtl", fontName=font_name, fontSize=10, leading=15, alignment=TA_RIGHT)
    heading = ParagraphStyle("rtl-title", parent=right, fontSize=17, leading=22, spaceAfter=6)
    story = [
        Paragraph(_rtl(company_name), heading),
        Paragraph(_rtl(title), heading),
        Paragraph(_rtl(subtitle), right),
    ]
    for line in metadata:
        story.append(Paragraph(_rtl(line), right))
    story.append(Spacer(1, 8))
    table_data = [[_rtl(value) for value in columns]] + [[_rtl(value) for value in row] for row in rows]
    table = Table(table_data, repeatRows=1, hAlign="RIGHT")
    table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f4e78")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("ALIGN", (0, 0), (-1, -1), "RIGHT"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#b7c9d6")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f3f7fa")]),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(table)
    doc.build(story)
    output.seek(0)
    return output


def build_docx(title: str, company_name: str, subtitle: str, columns: Sequence[str], rows: Iterable[Sequence[object]], metadata: Sequence[str]) -> BytesIO:
    output = BytesIO()
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    for text, size, bold in ((company_name, 18, True), (title, 16, True), (subtitle, 10, False)):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
    for line in metadata:
        paragraph = doc.add_paragraph(line)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for idx, value in enumerate(columns):
        cell = table.rows[0].cells[idx]
        cell.text = str(value)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = "" if value is None else str(value)
            cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.save(output)
    output.seek(0)
    return output


def build_excel(title: str, company_name: str, subtitle: str, columns: Sequence[str], rows: Iterable[Sequence[object]], metadata: Sequence[str]) -> BytesIO:
    output = BytesIO()
    wb = Workbook()
    ws = wb.active
    ws.title = "التقرير"
    ws.sheet_view.rightToLeft = True
    max_col = max(1, len(columns))
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max_col)
    ws["A1"] = company_name
    ws["A1"].font = Font(size=18, bold=True)
    ws["A1"].alignment = Alignment(horizontal="right")
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=max_col)
    ws["A2"] = title
    ws["A2"].font = Font(size=15, bold=True)
    ws["A2"].alignment = Alignment(horizontal="right")
    ws.merge_cells(start_row=3, start_column=1, end_row=3, end_column=max_col)
    ws["A3"] = subtitle
    ws["A3"].alignment = Alignment(horizontal="right")
    row_cursor = 4
    for line in metadata:
        ws.merge_cells(start_row=row_cursor, start_column=1, end_row=row_cursor, end_column=max_col)
        ws.cell(row_cursor, 1, line).alignment = Alignment(horizontal="right")
        row_cursor += 1
    row_cursor += 1
    header_row = row_cursor
    fill = PatternFill("solid", fgColor="1F4E78")
    thin = Side(style="thin", color="B7C9D6")
    for col_idx, value in enumerate(columns, 1):
        cell = ws.cell(header_row, col_idx, value)
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = fill
        cell.alignment = Alignment(horizontal="center")
        cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
    data_rows = list(rows)
    for r_idx, row in enumerate(data_rows, header_row + 1):
        for c_idx, value in enumerate(row, 1):
            cell = ws.cell(r_idx, c_idx, value)
            cell.alignment = Alignment(horizontal="right")
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            if isinstance(value, (int, float)):
                cell.number_format = '#,##0.00'
    ws.freeze_panes = f"A{header_row + 1}"
    ws.auto_filter.ref = f"A{header_row}:{get_column_letter(max_col)}{header_row + len(data_rows)}"
    for col_idx in range(1, max_col + 1):
        values = [str(ws.cell(r, col_idx).value or "") for r in range(header_row, header_row + len(data_rows) + 1)]
        ws.column_dimensions[get_column_letter(col_idx)].width = min(max(max(map(len, values)), 10) + 2, 34)
    wb.save(output)
    output.seek(0)
    return output
